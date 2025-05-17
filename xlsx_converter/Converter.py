import pandas as pd
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from fpdf import FPDF
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# TODO: nie działają polskie znaki w pdf...
# pewnie trzeba zmienić fpdf na fpdf2

# Custom PDF class with optional page numbering
class CustomPDF(FPDF):
    def __init__(self, add_page_numbers=False):
        super().__init__()
        self.add_page_numbers = add_page_numbers
        
    def footer(self):
        if self.add_page_numbers:
            self.set_y(-15)  # Position 15 mm from bottom
            self.set_font('Arial', 'I', 8)
            # Page number centered
            self.cell(0, 10, f'Strona {self.page_no()}', 0, 0, 'C')

class Converter:
    def __init__(self, path=""):
        self.path = path
        self.columns = []

    def open_file(self):
        df = pd.read_excel(self.path)
        return df

    def adjust_columns(self):
        df = self.open_file()
        self.columns = df.columns.to_list()
        return df

    def remove_empty_lines(self, text):
        lines = text.split('\n')
        non_empty_lines = [line for line in lines if line.strip() != '']
        return '\n'.join(non_empty_lines)

    def convert_into_pdf(self, font_size=10, interval=20, title="", file_name="mygfg.pdf",
                         selected_columns=None, single_line=True, alignment="left",
                         add_title_page=False, description="", add_page_numbers=False):
        df = self.adjust_columns()
        if selected_columns:
            df = df[selected_columns]
            
        # Use CustomPDF with page numbering option
        pdf = CustomPDF(add_page_numbers=add_page_numbers)
        pdf.set_font("Arial", size=font_size)

        if add_title_page:
            pdf.add_page()
            if title:
                pdf.set_font("Arial", size=font_size + 4)  # Slightly larger font for the title
                pdf.cell(0, 10, title.encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C')
                pdf.set_font("Arial", size=font_size)  # Reset font size
            if description:
                pdf.multi_cell(0, interval, description.encode('latin-1', 'replace').decode('latin-1'), align='L')
        else:
            if title or description:
                pdf.add_page()
                if title:
                    pdf.set_font("Arial", size=font_size + 4)  # Slightly larger font for the title
                    pdf.cell(0, 10, title.encode('latin-1', 'replace').decode('latin-1'), ln=True, align='C')
                    pdf.set_font("Arial", size=font_size)  # Reset font size
                if description:
                    pdf.multi_cell(0, interval, description.encode('latin-1', 'replace').decode('latin-1'), align='L')

        align_map = {"left": 'L', "center": 'C', "right": 'R'}

        for index, row in df.iterrows():
            if index == 0 and not add_title_page:
                pdf.add_page()  # Add the first page only if there's no title page
            elif index > 0:
                pdf.add_page()  # Add a new page for each row except the first one if there's no title page
            for col, value in row.items():
                if single_line:
                    pdf.set_font("Arial", 'B', size=font_size)  # Bold font for column names
                    col_width = pdf.get_string_width(f"{col}: ") + 2
                    x_before = pdf.get_x()
                    y_before = pdf.get_y()
                    pdf.multi_cell(col_width, interval, f"{col}: ".encode('latin-1', 'replace').decode('latin-1'),
                                   align=align_map[alignment])
                    pdf.set_xy(x_before + col_width, y_before)
                    pdf.set_font("Arial", size=font_size)  # Regular font for content
                    pdf.multi_cell(0, interval, f"{value}".encode('latin-1', 'replace').decode('latin-1'),
                                   align=align_map[alignment])
                else:
                    pdf.set_font("Arial", 'B', size=font_size)  # Bold font for column names
                    pdf.multi_cell(0, interval, f"{col}:".encode('latin-1', 'replace').decode('latin-1'),
                                   align=align_map[alignment])
                    pdf.set_font("Arial", size=font_size)  # Regular font for other text
                    pdf.multi_cell(0, interval, str(value).encode('latin-1', 'replace').decode('latin-1'),
                                   align=align_map[alignment])
        pdf.output(file_name)

    # Helper function to add page numbers in Word documents
    def _add_page_number(self, doc):
        sections = doc.sections
        for section in sections:
            footer = section.footer
            paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add page number field
            run = paragraph.add_run()
            fld_char1 = OxmlElement('w:fldChar')
            fld_char1.set(qn('w:fldCharType'), 'begin')
            run._r.append(fld_char1)
            
            instr_text = OxmlElement('w:instrText')
            instr_text.set(qn('xml:space'), 'preserve')
            instr_text.text = "PAGE"
            run._r.append(instr_text)
            
            fld_char2 = OxmlElement('w:fldChar')
            fld_char2.set(qn('w:fldCharType'), 'end')
            run._r.append(fld_char2)

    def convert_into_word(self, font_size=10, title="", file_name="TextToWord.docx",
                          selected_columns=None, line_spacing=1.0, single_line=True, alignment="left",
                          add_title_page=False, description="", add_page_numbers=False):
        df = self.adjust_columns()
        if selected_columns:
            df = df[selected_columns]
        doc = Document()
        if add_title_page:
            if title:
                title_paragraph = doc.add_heading(level=1)
                title_run = title_paragraph.add_run(title)
                title_run.font.size = Pt(font_size + 6)  # Larger font for the title
            if description:
                description_paragraph = doc.add_paragraph()
                description_run = description_paragraph.add_run(description)
                description_run.font.size = Pt(font_size + 2)  # Larger font for the description
        else:
            if title:
                title_paragraph = doc.add_heading(level=1)
                title_run = title_paragraph.add_run(title)
                title_run.font.size = Pt(font_size + 6)  # Larger font for the title
            if description:
                description_paragraph = doc.add_paragraph()
                description_run = description_paragraph.add_run(description)
                description_run.font.size = Pt(font_size + 2)  # Larger font for the description

        align_map = {"left": WD_ALIGN_PARAGRAPH.LEFT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                     "right": WD_ALIGN_PARAGRAPH.RIGHT}

        for index, row in df.iterrows():
            if index > 0 or (index == 0 and add_title_page):
                doc.add_page_break()  # Add a new page for each row except the first one if there's no title page
            for col, value in row.items():
                paragraph = doc.add_paragraph()
                paragraph.alignment = align_map[alignment]
                run = paragraph.add_run(f"{col}: ")
                run.bold = True  # Bold font for column names
                run.font.size = Pt(font_size)
                if single_line:
                    run = paragraph.add_run(f" {value}")
                else:
                    run = paragraph.add_run()
                    run.add_break()
                    run = paragraph.add_run(str(value))
                run.font.size = Pt(font_size)
                paragraph_format = paragraph.paragraph_format
                paragraph_format.line_spacing = line_spacing  # Set line spacing
        # Add page numbering only if requested
        if add_page_numbers:
            self._add_page_number(doc)
            
        doc.save(file_name)
